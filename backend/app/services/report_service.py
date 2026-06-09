"""Medical report upload, multi-format text extraction, and AI analysis."""
from __future__ import annotations

import asyncio
import csv
import hashlib
import io
import logging
import re
from pathlib import Path
from uuid import UUID, uuid4

import pdfplumber
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_settings
from app.models import Report

logger = logging.getLogger(__name__)
settings = get_settings()

UPLOAD_ROOT = Path(__file__).resolve().parent.parent.parent / "data" / "reports"
MAX_REPORT_BYTES = 15 * 1024 * 1024

SUPPORTED_FORMATS_LABEL = (
    "PDF, images (PNG/JPG/WebP/GIF/BMP/TIFF), Word (.docx), Excel (.xlsx), "
    "and text files (.txt, .csv, .json, .xml, .html, .md)"
)

EXTENSION_MIME: dict[str, str] = {
    ".pdf": "application/pdf",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
    ".bmp": "image/bmp",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".txt": "text/plain",
    ".csv": "text/csv",
    ".tsv": "text/tab-separated-values",
    ".md": "text/markdown",
    ".json": "application/json",
    ".xml": "application/xml",
    ".html": "text/html",
    ".htm": "text/html",
    ".log": "text/plain",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"}
TEXT_EXTENSIONS = {".txt", ".csv", ".tsv", ".md", ".json", ".xml", ".html", ".htm", ".log"}

_ROW_RE = re.compile(
    r"^(?P<test>[A-Za-z][A-Za-z0-9\s()/-]+?)\s+"
    r"(?P<value>[\d,.]+)\s*"
    r"(?:(?P<unit>[^\d\n]*?)\s+)?"
    r"(?P<low>[\d,.]+)\s*[-–]\s*(?P<high>[\d,.]+)",
    re.I,
)


def _file_path_for_key(s3_key: str) -> Path:
    key = s3_key.removeprefix("reports/")
    return UPLOAD_ROOT / key


def normalize_extension(filename: str, mime_type: str | None = None) -> str:
    ext = Path(filename or "").suffix.lower()
    if ext in EXTENSION_MIME:
        return ext
    mime = (mime_type or "").lower().split(";")[0].strip()
    for candidate, candidate_mime in EXTENSION_MIME.items():
        if mime == candidate_mime:
            return candidate
    if mime.startswith("image/"):
        return ".png"
    if mime == "text/plain":
        return ".txt"
    return ext or ".bin"


def validate_report_file(filename: str, mime_type: str | None, data: bytes) -> tuple[str, str]:
    if len(data) > MAX_REPORT_BYTES:
        raise ValueError("File is too large. Maximum size is 15 MB.")
    if not data:
        raise ValueError("Uploaded file is empty.")

    ext = normalize_extension(filename, mime_type)
    if ext == ".doc":
        raise ValueError(
            f"Legacy .doc files are not supported. Save as .docx or PDF. Supported: {SUPPORTED_FORMATS_LABEL}"
        )
    if ext == ".xls":
        raise ValueError(
            f"Legacy .xls files are not supported. Save as .xlsx or CSV. Supported: {SUPPORTED_FORMATS_LABEL}"
        )
    if ext == ".bin" or ext not in EXTENSION_MIME:
        raise ValueError(f"Unsupported file type. Supported: {SUPPORTED_FORMATS_LABEL}")

    mime = mime_type or EXTENSION_MIME[ext]
    return ext, mime


def save_report_file(patient_id: UUID, report_id: UUID, data: bytes, extension: str) -> str:
    ext = extension if extension.startswith(".") else f".{extension}"
    rel = f"{patient_id}/{report_id}{ext}"
    path = UPLOAD_ROOT / rel
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)
    return f"reports/{rel}"


def load_report_bytes(s3_key: str) -> bytes | None:
    path = _file_path_for_key(s3_key)
    if not path.is_file():
        return None
    return path.read_bytes()


def file_checksum(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def extract_text_from_pdf(data: bytes) -> str:
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())
    return "\n".join(parts).strip()


def extract_text_from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append("\t".join(cells))
    return "\n".join(chunks).strip()


def extract_text_from_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if values:
                lines.append("\t".join(values))
    workbook.close()
    return "\n".join(lines).strip()


def extract_text_from_plain(data: bytes, extension: str) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            text = ""
    else:
        return ""

    if extension == ".csv":
        reader = csv.reader(io.StringIO(text))
        return "\n".join("\t".join(row) for row in reader if any(cell.strip() for cell in row)).strip()
    if extension == ".tsv":
        reader = csv.reader(io.StringIO(text), delimiter="\t")
        return "\n".join("\t".join(row) for row in reader if any(cell.strip() for cell in row)).strip()
    return text.strip()


async def extract_text_from_image(data: bytes, mime_type: str) -> str:
    if not settings.gemini_api_key:
        raise ValueError(
            "Image and scan uploads need GEMINI_API_KEY configured for text extraction. "
            "Alternatively upload PDF, Word, Excel, or plain text."
        )

    import google.generativeai as genai

    def _run() -> str:
        genai.configure(api_key=settings.gemini_api_key)
        model = genai.GenerativeModel(settings.gemini_model)
        response = model.generate_content(
            [
                (
                    "Extract all readable text from this medical or lab report file. "
                    "Return plain text only with test names, values, units, and reference ranges."
                ),
                {"mime_type": mime_type, "data": data},
            ]
        )
        return (response.text or "").strip()

    try:
        return await asyncio.to_thread(_run)
    except Exception as exc:
        logger.warning("Image text extraction failed: %s", exc)
        raise ValueError(
            "Could not read text from this image. Try a clearer photo or upload a PDF/text report."
        ) from exc


async def extract_report_text(data: bytes, filename: str, mime_type: str | None = None) -> str:
    extension, resolved_mime = validate_report_file(filename, mime_type, data)

    if extension == ".pdf":
        text = extract_text_from_pdf(data)
    elif extension == ".docx":
        text = extract_text_from_docx(data)
    elif extension == ".xlsx":
        text = extract_text_from_xlsx(data)
    elif extension in TEXT_EXTENSIONS:
        text = extract_text_from_plain(data, extension)
    elif extension in IMAGE_EXTENSIONS:
        text = await extract_text_from_image(data, resolved_mime)
    else:
        text = extract_text_from_plain(data, extension)

    if not text.strip():
        raise ValueError(
            f"Could not extract readable text from this file. Supported: {SUPPORTED_FORMATS_LABEL}"
        )
    return text.strip()


def _parse_float(raw: str) -> float | None:
    try:
        return float(raw.replace(",", ""))
    except ValueError:
        return None


def _heuristic_analysis(ocr: str) -> dict:
    text = ocr.strip()
    if not text:
        return {
            "abnormal": [],
            "summary": "No readable text was found in the uploaded file. Try a clearer scan or a text-based document.",
        }

    lower = text.lower()
    if any(
        phrase in lower
        for phrase in (
            "within typical reference",
            "within reference range",
            "within normal",
            "all sample values",
            "within typical reference ranges",
        )
    ):
        return {
            "abnormal": [],
            "summary": (
                "Based on the extracted report text, values appear within typical reference ranges. "
                "Please discuss the full report with your physician."
            ),
        }

    abnormal: list[dict] = []
    seen_tests: set[str] = set()

    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        flag_match = re.search(r"\b(LOW|HIGH)\b", stripped, re.I)
        if flag_match:
            flag = flag_match.group(1).upper()
            before = stripped[: flag_match.start()].strip()
            value_match = re.search(r"([\d,.]+)\s*(?:g/dL|mg/dL|/μL|million)?", before, re.I)
            test = re.sub(r"[:]\s*$", "", before.split(value_match.group(0))[0].strip() if value_match else before)
            value = value_match.group(1) if value_match else "—"
            key = test.lower()
            if test and key not in seen_tests:
                seen_tests.add(key)
                abnormal.append({"test": test, "value": value, "flag": flag})
            continue

        row = _ROW_RE.match(stripped)
        if not row:
            continue
        test = row.group("test").strip()
        value_raw = row.group("value")
        low = _parse_float(row.group("low"))
        high = _parse_float(row.group("high"))
        value = _parse_float(value_raw)
        if low is None or high is None or value is None:
            continue
        flag = "NORMAL"
        if value < low:
            flag = "LOW"
        elif value > high:
            flag = "HIGH"
        key = test.lower()
        if flag != "NORMAL" and key not in seen_tests:
            seen_tests.add(key)
            abnormal.append({"test": test, "value": value_raw, "flag": flag})

    if abnormal:
        highlights = ", ".join(f"{a['test']} ({a['flag'].lower()})" for a in abnormal[:4])
        summary = (
            f"Some markers are outside the reference range: {highlights}. "
            "Please review these findings with your physician."
        )
    else:
        summary = (
            "No clear out-of-range values were detected in the extracted report text. "
            "Your doctor can interpret the full results in clinical context."
        )

    return {"abnormal": abnormal, "summary": summary}


async def _llm_analysis(ocr: str) -> dict | None:
    from app.multi_agent.llm import llm

    if not llm.available:
        return None

    prompt = (
        "Analyze this lab report text for a patient-facing healthcare assistant.\n"
        "Use ONLY values present in the report text. Do not invent results.\n"
        "Return ONLY JSON with this shape:\n"
        '{"abnormal": [{"test": "name", "value": "as shown", "flag": "LOW|HIGH|NORMAL"}], '
        '"summary": "2-3 sentence plain English explanation. No diagnosis or prescribing."}\n'
        "Include abnormal entries only for LOW or HIGH flags. Omit normal tests from abnormal list.\n\n"
        f"REPORT TEXT:\n{ocr[:12000]}"
    )
    parsed = await llm.json_prompt(prompt)
    if not parsed or not isinstance(parsed, dict):
        return None

    abnormal = parsed.get("abnormal") or []
    summary = (parsed.get("summary") or "").strip()
    if not summary:
        return None

    cleaned_abnormal = []
    for item in abnormal:
        if not isinstance(item, dict):
            continue
        flag = str(item.get("flag") or "").upper()
        if flag not in ("LOW", "HIGH"):
            continue
        test = str(item.get("test") or "").strip()
        if not test:
            continue
        cleaned_abnormal.append(
            {
                "test": test,
                "value": str(item.get("value") or "—"),
                "flag": flag,
            }
        )

    return {"abnormal": cleaned_abnormal, "summary": summary}


async def analyze_ocr_text(ocr: str) -> dict:
    llm_result = await _llm_analysis(ocr)
    if llm_result:
        return llm_result
    return _heuristic_analysis(ocr)


async def ensure_report_ocr(report: Report, *, refresh: bool = False) -> str:
    meta = (report.analysis_json or {}).get("_meta") or {}
    filename = meta.get("filename") or "medical-report.pdf"
    mime = report.mime_type or EXTENSION_MIME.get(normalize_extension(filename), "application/octet-stream")

    raw = load_report_bytes(report.s3_key)
    if raw:
        ocr = await extract_report_text(raw, filename, mime)
        if ocr:
            report.ocr_text = ocr
            return ocr

    if not refresh and report.ocr_text and report.ocr_text.strip():
        return report.ocr_text.strip()

    if not raw:
        raise ValueError("Report file not found on server. Please upload the file again.")

    raise ValueError(
        f"Could not extract text from this file. Supported: {SUPPORTED_FORMATS_LABEL}"
    )


async def analyze_report_record(
    db: AsyncSession,
    report: Report,
    *,
    force: bool = False,
) -> dict:
    meta = (report.analysis_json or {}).get("_meta") or {}
    if (
        not force
        and report.analysis_json
        and (report.analysis_json.get("summary") or report.analysis_json.get("abnormal") is not None)
        and report.ocr_text
    ):
        return report.analysis_json

    ocr = await ensure_report_ocr(report, refresh=force)
    analysis = await analyze_ocr_text(ocr)
    report.analysis_json = {**analysis, "_meta": meta}
    await db.flush()
    return report.analysis_json


async def create_and_analyze_report(
    db: AsyncSession,
    patient_id: UUID,
    data: bytes,
    filename: str,
    mime_type: str | None = None,
) -> Report:
    extension, resolved_mime = validate_report_file(filename, mime_type, data)
    await extract_report_text(data, filename, resolved_mime)

    report_id = uuid4()
    s3_key = save_report_file(patient_id, report_id, data, extension)
    report = Report(
        id=report_id,
        patient_id=patient_id,
        s3_key=s3_key,
        mime_type=resolved_mime,
        file_checksum=file_checksum(data),
        analysis_json={"_meta": {"filename": filename, "extension": extension}},
    )
    db.add(report)
    await db.flush()
    await analyze_report_record(db, report, force=True)
    return report
